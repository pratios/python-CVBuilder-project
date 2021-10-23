from docx import Document
from docx.shared import Inches
import pyttsx3

pyttsx3.speak('ah ah ah ah ah ah yes give it to me daddy yes yes yes fuck me so hard ah ah ah make me cum yes yes yes yes yes  I\'m cumming')

document = Document()

#Profile picture
document.add_picture('LinkedInProfile.jpg', width = Inches(2.0))

name = input('What is your name?')
phone = input('What is your phone number?')
email = input('What is your email id?')

#Contact details
document.add_paragraph(
    name + ' | '+phone+' | '+email
)

#About me
document.add_heading('About me')
document.add_paragraph(
    input('Tell me about yourself : ')
)

#Work experience
document.add_heading('Work Experience')
p = document.add_paragraph()

company = input('Company : ')
from_date = input('From : ')
to_date = input('To : ')

p.add_run(company + ' ').bold = True
p.add_run(from_date + '-'+to_date+'\n').italic = True

experience_details = input('Describe your company experience at '+company+' : ')
p.add_run(experience_details)

#more experiences
while True :
    has_more_experiences = input('Do you have more experience? (y/n)')

    if has_more_experiences.lower() == 'y' :
        
        company = input('Company : ')
        from_date = input('From : ')
        to_date = input('To : ')

        p.add_run('\n'+company + ' ').bold = True
        p.add_run(from_date + '-'+to_date+'\n').italic = True

        experience_details = input('Describe your company experience at '+company+' : ')
        p.add_run(experience_details)
    else:
        break;

#Skills
document.add_heading('Skills')

skill = input('Enter your skill : ')
p = document.add_paragraph(skill)
p.style = 'List Bullet'

while True : 
     more_skills = input('Do you want to enter more skills? y/n : ')

     if more_skills.lower() == 'y' :
         skill = input('Enter your skill : ')
         p = document.add_paragraph(skill)
         p.style = 'List Bullet'
     else :
        break

#Footer
section = document.sections[0]
footer = section.footer
p = footer.paragraphs[0]
p.text = 'CV generated using CV builder.'

document.save('CV.docx')